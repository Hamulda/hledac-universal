"""
Document Metadata Extractor - FOCA-style forensic metadata extraction
=================================================================

FOCA (Fingerprinting Organizations with Collected Archives) style extraction:
- Author, last_modified_by, company, template path -> identity & infrastructure
- Embedded fonts (foreign/specific = geographic indication)
- Revision history in DOCX (tracking deleted text in revisions)
- GPS coordinates in EXIF (images embedded in documents)
- PDF hidden content: invisible layers, form fields, JavaScript, embedded files
- Office document macros analysis (olevba integration)
- Email header forensics
- Presentation forensics (speaker notes, hidden slides)
- CAD/technical drawing metadata (DXF, DWG, SVG)

M1 8GB RAM optimized - bounded extraction with fail-safe throughout.

Sprint 52: Document Metadata Extractor
Extended: Sprint FOCADI-16 FOCA-style metadata pipeline
"""

from __future__ import annotations

import asyncio
import hashlib
import io
import logging
import re
import sqlite3
import zipfile
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)

# =============================================================================
# LIBRARY AVAILABILITY FLAGS
# =============================================================================

FITZ_AVAILABLE = False
DOCX_AVAILABLE = False
OPENPYXL_AVAILABLE = False
PIL_AVAILABLE = False
OLEVB_AVAILABLE = False

try:
    import fitz  # PyMuPDF
    FITZ_AVAILABLE = True
except ImportError:
    pass

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    pass

try:
    import openpyxl
    OPENPYXL_AVAILABLE = True
except ImportError:
    pass

try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    pass

try:
    import olevba
    OLEVB_AVAILABLE = True
except ImportError:
    pass

# =============================================================================
# CONSTANTS
# =============================================================================

SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.xlsx', '.pptx', '.odt', '.svg', '.dxf', '.eml', '.msg'}
MAX_INTERNAL_PATHS = 100
MAX_EMBEDDED_FONTS = 50
MAX_PDF_OBJECTS = 500
MAX_MACRO_ANALYSIS_CHARS = 100_000
MAX_EMAIL_HEADERS = 100
MAX_REVISIONS = 50
MAX_HIDDEN_SLIDES = 20

# Sprint F179B: canonical cache root via paths.py (M1/RAMDISK-safe)
try:
    from hledac.universal.paths import CACHE_ROOT
except ImportError:
    from pathlib import Path
    CACHE_ROOT = Path("/tmp/hledac_cache")

MAX_GPS_COORDS = 20
CACHE_TTL_DAYS = 30
EXTRACTION_TIMEOUT = 10.0
PDF_DEEP_TIMEOUT = 15.0

# Cache DB path - under RAMDISK CACHE_ROOT, not home-relative
CACHE_DIR = CACHE_ROOT
CACHE_DB_PATH = CACHE_DIR / 'doc_meta_cache.db'

# Regex patterns for internal paths
INTERNAL_PATH_PATTERNS = {
    'windows': re.compile(r'[A-Za-z]:\\[^<>"\'\s]{3,}'),
    'unc': re.compile(r'\\\\[^<>"\'\s]{3,}'),
    'unix': re.compile(r'(?:/home/|/Users/|/var/|/etc/)[^<>"\'\s]{2,}'),
}

# Macro detection
VBA_PROJECT_PATTERNS = [
    b'vbaProject.bin',
    b'xl/vbaProject.bin',
    b'word/vbaProject.bin',
]

# PDF macro detection patterns
PDF_MACRO_PATTERNS = [
    b'/JS',
    b'/JavaScript',
    b'/Launch',
]

# C2 URL patterns in macros (common malware patterns)
C2_URL_PATTERNS = [
    re.compile(r'https?://[^\s<>"\']+\.(php|asp|jsp|cgi|pl)', re.IGNORECASE),
    re.compile(r'https?://\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}', re.IGNORECASE),
    re.compile(r'https?://[^\s<>"\']+\.tk\.', re.IGNORECASE),
    re.compile(r'https?://[^\s<>"\']+\.ga\.', re.IGNORECASE),
    re.compile(r'https?://[^\s<>"\']+\.cf\.', re.IGNORECASE),
    re.compile(r'https?://[^\s<>"\']+\.pw\.', re.IGNORECASE),
]

# Suspicious API call patterns in macros
SUSPICIOUS_API_PATTERNS = [
    re.compile(r'CreateObject|Wscript\.Shell', re.IGNORECASE),
    re.compile(r'UrlDownloadToFile|XMLHTTP', re.IGNORECASE),
    re.compile(r'WinHttp\.WinHttpRequest', re.IGNORECASE),
    re.compile(r'Process32|Module32', re.IGNORECASE),
    re.compile(r'RegOpenKey|RegSetValue', re.IGNORECASE),
]

# Email header analysis patterns
RECEIVED_HEADER_PATTERN = re.compile(r'from\s+([^\s\(]+)\s+\(?(.*?)\)?\s+by\s+([^\s\(]+)', re.IGNORECASE)
X_ORIGINATING_IP_PATTERN = re.compile(r'X-Originating-IP:\s*\[?(\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3})\]?', re.IGNORECASE)
MESSAGE_ID_DOMAIN_PATTERN = re.compile(r'@([a-zA-Z0-9\-\.]+)>')

# SVG metadata patterns
SVG_METADATA_PATTERNS = {
    'author': re.compile(r'<dc:creator[^>]*>([^<]+)</dc:creator>', re.IGNORECASE),
    'title': re.compile(r'<dc:title[^>]*>([^<]+)</dc:title>', re.IGNORECASE),
    'description': re.compile(r'<dc:description[^>]*>([^<]+)</dc:description>', re.IGNORECASE),
}


class _DocumentMetadataExtractor:
    """
    FOCA-style forensic metadata extractor for documents.

    Extracts:
    - Author, creator, organization, company, template path
    - Internal file paths (Windows/UNC/Unix)
    - GPS coordinates from embedded images
    - Revision history (DOCX)
    - Embedded fonts (geographic indication)
    - PDF hidden content (invisible layers, forms, JS, embedded files)
    - Macro analysis (olevba integration)
    - Email header forensics
    - Presentation forensics (speaker notes, hidden slides)
    - CAD/SVG metadata

    CPU-heavy operations run in executor with timeout.
    Results cached in SQLite with 30-day TTL.
    """

    def __init__(self):
        """Initialize extractor with SQLite cache."""
        self._init_cache()
        logger.debug("[DOCMETA] FOCA DocumentMetadataExtractor initialized")

    def _init_cache(self) -> None:
        """Initialize SQLite cache."""
        try:
            CACHE_DIR.mkdir(parents=True, exist_ok=True)
            self._conn = sqlite3.connect(str(CACHE_DB_PATH), timeout=5.0)
            self._conn.execute('''
                CREATE TABLE IF NOT EXISTS doc_meta_cache (
                    key TEXT PRIMARY KEY,
                    value TEXT NOT NULL,
                    timestamp INTEGER NOT NULL
                )
            ''')
            self._conn.execute('''
                CREATE INDEX IF NOT EXISTS idx_timestamp
                ON doc_meta_cache(timestamp)
            ''')
            self._conn.commit()
        except Exception as e:
            logger.warning(f"[DOCMETA] Cache init failed: {e}")
            self._conn = None

    def _get_cache_key(self, content: bytes) -> str:
        """Generate cache key from first 1024 bytes."""
        prefix = content[:1024]
        return hashlib.sha256(prefix).hexdigest()

    def _is_cache_valid(self, timestamp: int) -> bool:
        """Check if cache entry is still valid (within TTL)."""
        import time
        age_days = (time.time() - timestamp) / 86400
        return age_days < CACHE_TTL_DAYS

    def _get_cached(self, content: bytes) -> dict | None:
        """Get cached extraction result."""
        if not self._conn:
            return None
        try:
            key = self._get_cache_key(content)
            cursor = self._conn.execute(
                'SELECT value, timestamp FROM doc_meta_cache WHERE key = ?',
                (key,)
            )
            row = cursor.fetchone()
            if row and self._is_cache_valid(row[1]):
                import json
                return json.loads(row[0])
        except Exception:
            pass
        return None

    def _cache(self, content: bytes, result: dict) -> None:
        """Cache extraction result."""
        if not self._conn or not result:
            return
        try:
            import json
            import time
            key = self._get_cache_key(content)
            value = json.dumps(result)
            timestamp = int(time.time())
            self._conn.execute(
                'INSERT OR REPLACE INTO doc_meta_cache (key, value, timestamp) VALUES (?, ?, ?)',
                (key, value, timestamp)
            )
            self._conn.commit()
        except Exception:
            pass

    def _get_extension(self, url: str) -> str:
        """Get file extension from URL."""
        path = url.lower().rsplit('/', 1)[-1] if '/' in url else url
        ext = '.' + path.rsplit('.', 1)[-1] if '.' in path else ''
        return ext if ext in SUPPORTED_EXTENSIONS else ''

    async def extract(self, content: bytes, url: str) -> dict:
        """
        Extract FOCA-style forensic metadata from document.

        Args:
            content: Raw document bytes
            url: Source URL for extension detection

        Returns:
            Dict with keys: author, creator, organization, company, template_path,
            last_modified_by, revision_count, internal_paths, gps_coords,
            has_macros, macro_analysis, embedded_fonts, hidden_content,
            email_headers, presentation_notes, cad_metadata, format
        """
        ext = self._get_extension(url)
        if not ext or ext not in SUPPORTED_EXTENSIONS:
            return {}

        # Check cache
        cached = self._get_cached(content)
        if cached is not None:
            return cached

        try:
            loop = asyncio.get_running_loop()
            result = await asyncio.wait_for(
                loop.run_in_executor(None, self._extract_sync, content, ext),
                timeout=EXTRACTION_TIMEOUT
            )
            if result:
                self._cache(content, result)
            return result
        except TimeoutError:
            logger.debug(f"[DOCMETA] Timeout extracting from {url}")
            return {}
        except Exception as e:
            logger.debug(f"[DOCMETA] Extraction failed for {url}: {e}")
            return {}

    def _extract_sync(self, content: bytes, ext: str) -> dict:
        """Blocking extraction - runs in executor."""
        extractors = {
            '.pdf': self._extract_pdf,
            '.docx': self._extract_docx,
            '.xlsx': self._extract_xlsx,
            '.pptx': self._extract_pptx,
            '.odt': self._extract_odt,
            '.svg': self._extract_svg,
            '.dxf': self._extract_dxf,
            '.eml': self._extract_email,
            '.msg': self._extract_msg,
        }
        return extractors.get(ext, lambda _: {})(content)

    # =========================================================================
    # PDF EXTRACTION
    # =========================================================================

    def _extract_pdf(self, content: bytes) -> dict:
        """Extract from PDF using PyMuPDF with FOCA-style deep analysis."""
        result: dict[str, Any] = {'format': 'pdf'}

        if not FITZ_AVAILABLE:
            return self._extract_pdf_fallback(content, result)

        try:
            doc = fitz.open(stream=content, filetype='pdf')

            # Basic metadata
            meta = doc.metadata
            result['author'] = meta.get('author') or None
            result['creator'] = meta.get('creator') or None
            result['organization'] = meta.get('producer') or None
            result['last_modified_by'] = meta.get('modDate') or None
            result['title'] = meta.get('title') or None
            result['subject'] = meta.get('subject') or None

            # Company (often in producer field for MS Office PDFs)
            producer = meta.get('producer') or ''
            if 'microsoft' in producer.lower() or 'adobe' in producer.lower():
                result['company'] = self._extract_company_from_pdf_producer(producer)

            # Page count
            result['page_count'] = len(doc)

            # Revision count (from PDF info dictionary)
            result['revision_count'] = self._extract_pdf_revisions(doc, content)

            # Internal paths (search first 10 pages)
            internal_paths = []
            for page_num in range(min(10, len(doc))):
                page = doc[page_num]
                text = page.get_text()
                paths = self._find_internal_paths(text)
                internal_paths.extend(paths)
            result['internal_paths'] = list(set(internal_paths))[:MAX_INTERNAL_PATHS]

            # Embedded fonts (geographic indication)
            result['embedded_fonts'] = self._extract_pdf_fonts(doc)[:MAX_EMBEDDED_FONTS]

            # GPS coords from embedded images
            result['gps_coords'] = self._extract_pdf_gps(doc)

            # Hidden content analysis (PDF deep extraction)
            result['hidden_content'] = self._extract_pdf_hidden_content(doc, content)

            # Macro detection
            prefix = content[:50000]
            result['has_macros'] = any(p in prefix for p in PDF_MACRO_PATTERNS)

            # Template path (for Office-generated PDFs)
            result['template_path'] = self._extract_pdf_template_path(doc, content)

            doc.close()
            return result

        except Exception as e:
            logger.debug(f"[DOCMETA] PDF extraction failed: {e}")
            return self._extract_pdf_fallback(content, result)

    def _extract_pdf_fallback(self, content: bytes, result: dict) -> dict:
        """Fallback PDF extraction without PyMuPDF."""
        result['format'] = 'pdf'
        result['author'] = None
        result['creator'] = None
        result['organization'] = None
        result['last_modified_by'] = None
        result['title'] = None
        result['subject'] = None
        result['company'] = None
        result['page_count'] = 0
        result['revision_count'] = 0
        result['internal_paths'] = []
        result['embedded_fonts'] = []
        result['gps_coords'] = []
        result['hidden_content'] = {}
        result['template_path'] = None

        # Macro detection
        prefix = content[:50000]
        result['has_macros'] = any(p in prefix for p in PDF_MACRO_PATTERNS)

        return result

    def _extract_pdf_revisions(self, doc, content: bytes) -> int:
        """Extract revision count from PDF."""
        try:
            info = doc.metadata
            if info:
                for key, value in info.items():
                    if key and 'revision' in key.lower():
                        try:
                            return int(value)
                        except (ValueError, TypeError):
                            pass
            return 0
        except Exception:
            return 0

    def _extract_pdf_fonts(self, doc) -> list[dict[str, str]]:
        """Extract embedded fonts for geographic indication."""
        fonts = []
        try:
            for page_num in range(min(5, len(doc))):
                page = doc[page_num]
                font_list = page.get_fonts()
                for font in font_list[:10]:
                    try:
                        name = font[0] if font else None
                        ext = font[1] if len(font) > 1 else None
                        type_ = font[2] if len(font) > 2 else None
                        if name:
                            fonts.append({
                                'name': name,
                                'type': type_,
                                'embedding': ext
                            })
                    except Exception:
                        continue
        except Exception:
            pass
        return fonts

    def _extract_pdf_template_path(self, doc, content: bytes) -> str | None:
        """Extract template path from PDF metadata."""
        try:
            meta = doc.metadata
            if meta:
                for key, value in meta.items():
                    if value and ('template' in key.lower() or 'template' in str(value).lower()):
                        return str(value)
            return None
        except Exception:
            return None

    def _extract_pdf_hidden_content(self, doc, content: bytes) -> dict[str, Any]:
        """
        Extract PDF hidden content:
        - Invisible text layers (OCR vs embedded text mismatch)
        - Hidden form fields
        - JavaScript actions
        - Embedded files
        - Incremental updates
        """
        hidden: dict[str, Any] = {
            'has_invisible_text': False,
            'has_form_fields': False,
            'has_javascript': False,
            'embedded_files': [],
            'incremental_updates': 0,
            'suspicious_objects': [],
        }

        try:
            # Check for form fields
            for page_num in range(min(10, len(doc))):
                page = doc[page_num]
                widgets = page.widgets()
                if widgets:
                    hidden['has_form_fields'] = True

            # Check for JavaScript
            if doc.has_javascript():
                hidden['has_javascript'] = True

            # Check for embedded files
            try:
                embeds = doc.embfile_names
                for name in embeds[:20]:
                    hidden['embedded_files'].append(name)
            except Exception:
                pass

            # Count incremental updates (xref tables)
            xref_count = doc.xref_length()
            if xref_count > 10:
                hidden['incremental_updates'] = xref_count - 1

            # Scan for suspicious objects
            for i in range(min(MAX_PDF_OBJECTS, xref_count)):
                try:
                    obj_str = doc.xref_object(i, compressed=False)
                    if obj_str:
                        if any(p in obj_str for p in ['/JS', '/JavaScript', '/Launch']):
                            hidden['suspicious_objects'].append({
                                'xref': i,
                                'type': 'script_or_launch'
                            })
                except Exception:
                    continue

        except Exception as e:
            logger.debug(f"[DOCMETA] PDF hidden content extraction failed: {e}")

        return hidden

    def _extract_pdf_gps(self, doc) -> list[dict]:
        """Extract GPS coordinates from embedded images."""
        if not PIL_AVAILABLE:
            return []

        gps_coords = []
        try:
            for page_num in range(min(5, len(doc))):
                page = doc[page_num]
                images = page.get_images()
                for img in images:
                    try:
                        pix = fitz.Pixmap(doc, img[0])
                        if pix.n - pix.alpha > 3:
                            img_data = pix.tobytes('png')
                            img_obj = io.BytesIO(img_data)
                            img_obj = Image.open(img_obj)
                            exif = img_obj._getexif()
                            if exif:
                                gps = self._parse_exif_gps(exif)
                                if gps:
                                    gps_coords.append(gps)
                                    if len(gps_coords) >= MAX_GPS_COORDS:
                                        break
                    except Exception:
                        continue
                if len(gps_coords) >= MAX_GPS_COORDS:
                    break
        except Exception:
            pass

        return gps_coords[:MAX_GPS_COORDS]

    def _parse_exif_gps(self, exif) -> dict | None:
        """Parse GPS from EXIF data."""
        try:
            if 34853 not in exif:
                return None
            gps_ifd = exif[34853]

            lat = None
            lat_ref = None
            lon = None
            lon_ref = None

            if 1 in gps_ifd:
                lat_ref = gps_ifd[1]
            if 2 in gps_ifd:
                lat = _exif_to_float(gps_ifd[2])
            if 3 in gps_ifd:
                lon_ref = gps_ifd[3]
            if 4 in gps_ifd:
                lon = _exif_to_float(gps_ifd[4])

            if lat and lon:
                if lat_ref == 'S':
                    lat = -lat
                if lon_ref == 'W':
                    lon = -lon
                return {'lat': lat, 'lon': lon, 'source': 'exif'}

            return None
        except Exception:
            return None

    def _extract_company_from_pdf_producer(self, producer: str) -> str | None:
        """Extract company name from PDF producer string."""
        if not producer:
            return None
        patterns = [
            (r'Microsoft Office', 'Microsoft'),
            (r'Adobe Acrobat', 'Adobe'),
            (r'LibreOffice', 'LibreOffice'),
            (r'OpenOffice', 'Apache OpenOffice'),
            (r'Word', 'Microsoft'),
            (r'Excel', 'Microsoft'),
            (r'PowerPoint', 'Microsoft'),
        ]
        for pattern, company in patterns:
            if re.search(pattern, producer, re.IGNORECASE):
                return company
        return None

    # =========================================================================
    # DOCX EXTRACTION
    # =========================================================================

    def _extract_docx(self, content: bytes) -> dict:
        """Extract from DOCX with FOCA-style revision history and fonts."""
        result: dict[str, Any] = {'format': 'docx'}

        if not DOCX_AVAILABLE:
            return self._extract_docx_fallback(content, result)

        try:
            doc = docx.Document(io.BytesIO(content))

            # Core properties
            core = doc.core_properties
            result['author'] = core.author or None
            result['creator'] = core.last_modified_by or None
            result['last_modified_by'] = core.last_modified_by or None
            result['revision_count'] = core.revision or 0
            result['title'] = core.title or None
            result['subject'] = core.subject or None

            # Company and template path from custom properties
            result['company'] = None
            result['template_path'] = None

            # Extract from ZIP directly for extended properties
            result.update(self._extract_docx_extended_props(content))

            # Embedded fonts
            result['embedded_fonts'] = self._extract_docx_fonts(doc)[:MAX_EMBEDDED_FONTS]

            # Revision history (track changes)
            result['revision_history'] = self._extract_docx_revisions(content)[:MAX_REVISIONS]

            # Internal paths
            internal_paths = []
            for para in doc.paragraphs:
                paths = self._find_internal_paths(para.text)
                internal_paths.extend(paths)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        paths = self._find_internal_paths(cell.text)
                        internal_paths.extend(paths)
            result['internal_paths'] = list(set(internal_paths))[:MAX_INTERNAL_PATHS]

            # Macro detection via ZIP
            result['has_macros'] = self._check_docx_macros(content)

            # Macro analysis with olevba
            if result['has_macros'] and OLEVB_AVAILABLE:
                result['macro_analysis'] = self._analyze_macros_olevba(content, 'docm')

            result['gps_coords'] = []
            result['page_count'] = 0

            return result

        except Exception as e:
            logger.debug(f"[DOCMETA] DOCX extraction failed: {e}")
            return self._extract_docx_fallback(content, result)

    def _extract_docx_fallback(self, content: bytes, result: dict) -> dict:
        """Fallback DOCX extraction without python-docx."""
        result['format'] = 'docx'
        result['author'] = None
        result['creator'] = None
        result['last_modified_by'] = None
        result['organization'] = None
        result['company'] = None
        result['title'] = None
        result['subject'] = None
        result['template_path'] = None
        result['revision_count'] = 0
        result['revision_history'] = []
        result['embedded_fonts'] = []
        result['internal_paths'] = self._find_internal_paths(content.decode('utf-8', errors='ignore'))
        result['has_macros'] = self._check_docx_macros(content)
        result['macro_analysis'] = {}
        result['gps_coords'] = []
        result['page_count'] = 0
        return result

    def _extract_docx_extended_props(self, content: bytes) -> dict[str, Any]:
        """Extract extended properties from DOCX ZIP."""
        props = {}
        try:
            with zipfile.ZipFile(io.BytesIO(content)) as zf:
                if 'docProps/app.xml' in zf.namelist():
                    app_xml = zf.read('docProps/app.xml').decode('utf-8', errors='ignore')
                    company_match = re.search(r'<Company>([^<]*)</Company>', app_xml)
                    if company_match:
                        props['company'] = company_match.group(1)

                    template_match = re.search(r'<Template>([^<]*)</Template>', app_xml)
                    if template_match:
                        props['template_path'] = template_match.group(1)

                if 'docProps/custom.xml' in zf.namelist():
                    custom_xml = zf.read('docProps/custom.xml').decode('utf-8', errors='ignore')
                    company_matches = re.findall(r'<property[^>]*name="[^"]*company[^"]*"[^>]*><value>([^<]*)</value>', custom_xml, re.IGNORECASE)
                    if company_matches and 'company' not in props:
                        props['company'] = company_matches[0]

        except Exception as e:
            logger.debug(f"[DOCMETA] DOCX extended props extraction failed: {e}")
        return props

    def _extract_docx_fonts(self, doc) -> list[dict[str, str]]:
        """Extract embedded fonts from DOCX document."""
        fonts = []
        try:
            for para in doc.paragraphs:
                for run in para.runs:
                    if run.font:
                        font_info = {
                            'name': run.font.name,
                            'family': str(run.font.family) if run.font.family else None,
                        }
                        if font_info['name'] and font_info not in fonts:
                            fonts.append(font_info)
        except Exception:
            pass
        return fonts

    def _extract_docx_revisions(self, content: bytes) -> list[dict[str, Any]]:
        """Extract revision history from DOCX (track changes)."""
        revisions = []
        try:
            with zipfile.ZipFile(io.BytesIO(content)) as zf:
                if 'word/document.xml' in zf.namelist():
                    doc_xml = zf.read('word/document.xml').decode('utf-8', errors='ignore')

                    ins_pattern = re.compile(r'<w:ins[^>]*w:id="(\d+)"[^>]*w:author="([^"]*)"[^>]*w:date="([^"]*)"', re.IGNORECASE)
                    del_pattern = re.compile(r'<w:del[^>]*w:id="(\d+)"[^>]*w:author="([^"]*)"[^>]*w:date="([^"]*)"', re.IGNORECASE)

                    for match in ins_pattern.finditer(doc_xml):
                        revisions.append({
                            'id': match.group(1),
                            'author': match.group(2),
                            'date': match.group(3),
                            'type': 'insertion'
                        })

                    for match in del_pattern.finditer(doc_xml):
                        revisions.append({
                            'id': match.group(1),
                            'author': match.group(2),
                            'date': match.group(3),
                            'type': 'deletion'
                        })

        except Exception as e:
            logger.debug(f"[DOCMETA] DOCX revision extraction failed: {e}")
        return revisions[:MAX_REVISIONS]

    # =========================================================================
    # XLSX EXTRACTION
    # =========================================================================

    def _extract_xlsx(self, content: bytes) -> dict:
        """Extract from XLSX with embedded fonts."""
        result: dict[str, Any] = {'format': 'xlsx'}

        if not OPENPYXL_AVAILABLE:
            return self._extract_xlsx_fallback(content, result)

        try:
            wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)

            props = wb.properties
            result['author'] = props.creator or None
            result['creator'] = props.creator or None
            result['last_modified_by'] = props.lastModifiedBy or None
            result['organization'] = None
            result['title'] = props.title or None
            result['subject'] = props.subject or None

            result.update(self._extract_xlsx_extended_props(content))

            result['revision_count'] = 0

            internal_paths = []
            cell_count = 0
            for sheet in wb.worksheets:
                for row in sheet.iter_rows():
                    for cell in row:
                        if cell.value and isinstance(cell.value, str):
                            paths = self._find_internal_paths(cell.value)
                            internal_paths.extend(paths)
                        cell_count += 1
                        if cell_count >= 1000:
                            break
                    if cell_count >= 1000:
                        break
                if cell_count >= 1000:
                    break

            result['internal_paths'] = list(set(internal_paths))[:MAX_INTERNAL_PATHS]

            result['embedded_fonts'] = []

            result['has_macros'] = self._check_docx_macros(content)

            if result['has_macros'] and OLEVB_AVAILABLE:
                result['macro_analysis'] = self._analyze_macros_olevba(content, 'xlsm')

            result['gps_coords'] = []
            result['page_count'] = 0

            wb.close()
            return result

        except Exception as e:
            logger.debug(f"[DOCMETA] XLSX extraction failed: {e}")
            return self._extract_xlsx_fallback(content, result)

    def _extract_xlsx_extended_props(self, content: bytes) -> dict[str, Any]:
        """Extract extended properties from XLSX ZIP."""
        props = {}
        try:
            with zipfile.ZipFile(io.BytesIO(content)) as zf:
                if 'docProps/app.xml' in zf.namelist():
                    app_xml = zf.read('docProps/app.xml').decode('utf-8', errors='ignore')
                    company_match = re.search(r'<Company>([^<]*)</Company>', app_xml)
                    if company_match:
                        props['company'] = company_match.group(1)
        except Exception:
            pass
        return props

    def _extract_xlsx_fallback(self, content: bytes, result: dict) -> dict:
        """Fallback XLSX extraction without openpyxl."""
        result['format'] = 'xlsx'
        result['author'] = None
        result['creator'] = None
        result['last_modified_by'] = None
        result['organization'] = None
        result['company'] = None
        result['title'] = None
        result['subject'] = None
        result['template_path'] = None
        result['revision_count'] = 0
        result['embedded_fonts'] = []
        result['internal_paths'] = []
        result['has_macros'] = self._check_docx_macros(content)
        result['macro_analysis'] = {}
        result['gps_coords'] = []
        result['page_count'] = 0
        return result

    # =========================================================================
    # PPTX EXTRACTION
    # =========================================================================

    def _extract_pptx(self, content: bytes) -> dict:
        """Extract from PPTX with speaker notes and hidden slides."""
        result: dict[str, Any] = {'format': 'pptx'}

        try:
            with zipfile.ZipFile(io.BytesIO(content)) as zf:
                names = zf.namelist()

                if 'docProps/core.xml' in names:
                    core_xml = zf.read('docProps/core.xml').decode('utf-8', errors='ignore')
                    result['author'] = self._extract_xml_value(core_xml, 'dc:creator')
                    result['last_modified_by'] = self._extract_xml_value(core_xml, 'cp:lastModifiedBy')
                    result['title'] = self._extract_xml_value(core_xml, 'dc:title')
                    result['subject'] = self._extract_xml_value(core_xml, 'dc:subject')

                if 'docProps/app.xml' in names:
                    app_xml = zf.read('docProps/app.xml').decode('utf-8', errors='ignore')
                    result['company'] = self._extract_xml_value(app_xml, 'Company')
                    result['template_path'] = self._extract_xml_value(app_xml, 'Template')

                result['speaker_notes'] = self._extract_pptx_speaker_notes(zf, names)
                result['hidden_slides'] = self._extract_pptx_hidden_slides(zf, names)

                slide_count = len([n for n in names if re.match(r'ppt/slides/slide\d+\.xml', n)])
                result['slide_count'] = slide_count

                internal_paths = []
                for name in names:
                    if name.endswith('.xml'):
                        try:
                            xml_content = zf.read(name).decode('utf-8', errors='ignore')
                            paths = self._find_internal_paths(xml_content)
                            internal_paths.extend(paths)
                        except Exception:
                            pass
                result['internal_paths'] = list(set(internal_paths))[:MAX_INTERNAL_PATHS]

                result['has_macros'] = self._check_docx_macros(content)

                if result['has_macros'] and OLEVB_AVAILABLE:
                    result['macro_analysis'] = self._analyze_macros_olevba(content, 'pptm')

                result['gps_coords'] = []
                result['embedded_fonts'] = []

        except Exception as e:
            logger.debug(f"[DOCMETA] PPTX extraction failed: {e}")
            result = {'format': 'pptx', 'error': str(e)}

        return result

    def _extract_pptx_speaker_notes(self, zf, names: list[str]) -> list[str]:
        """Extract speaker notes from PPTX."""
        notes = []
        try:
            for name in names:
                if re.match(r'ppt/notesSlides/notesSlide\d+\.xml', name):
                    try:
                        xml_content = zf.read(name).decode('utf-8', errors='ignore')
                        text_matches = re.findall(r'<a:t>([^<]+)</a:t>', xml_content)
                        if text_matches:
                            note_text = ' '.join(text_matches)
                            if note_text.strip():
                                notes.append(note_text.strip())
                    except Exception:
                        pass
        except Exception:
            pass
        return notes[:50]

    def _extract_pptx_hidden_slides(self, zf, names: list[str]) -> list[dict[str, Any]]:
        """Extract hidden slides from PPTX."""
        hidden = []
        try:
            if 'ppt/presentation.xml' in names:
                pres_xml = zf.read('ppt/presentation.xml').decode('utf-8', errors='ignore')

                visible_ids = set()
                for sld_id in re.finditer(r'<p:sldId\b[^>]*>', pres_xml):
                    attrs = sld_id.group(0)
                    if 'show' not in attrs or 'show="1"' in attrs.lower():
                        rid_match = re.search(r'r:id="([^"]*)"', attrs)
                        if rid_match:
                            visible_ids.add(rid_match.group(1))

                if 'ppt/_rels/presentation.xml.rels' in names:
                    rels_xml = zf.read('ppt/_rels/presentation.xml.rels').decode('utf-8', errors='ignore')
                    for rel_match in re.finditer(r'<Relationship[^>]*Id="([^"]*)"[^>]*Target="([^"]*)"', rels_xml):
                        rid = rel_match.group(1)
                        target = rel_match.group(2)
                        if 'slides/slide' in target and rid not in visible_ids:
                            slide_num = re.search(r'slide(\d+)', target)
                            if slide_num:
                                hidden.append({
                                    'slide_number': int(slide_num.group(1)),
                                    'relationship_id': rid
                                })

        except Exception as e:
            logger.debug(f"[DOCMETA] PPTX hidden slides extraction failed: {e}")
        return hidden[:MAX_HIDDEN_SLIDES]

    # =========================================================================
    # ODT EXTRACTION
    # =========================================================================

    def _extract_odt(self, content: bytes) -> dict:
        """Extract from ODT (OpenDocument Text)."""
        result: dict[str, Any] = {'format': 'odt'}

        try:
            with zipfile.ZipFile(io.BytesIO(content)) as zf:
                names = zf.namelist()

                if 'meta.xml' in names:
                    meta_xml = zf.read('meta.xml').decode('utf-8', errors='ignore')
                    result['author'] = self._extract_xml_value(meta_xml, 'meta:initial-creator')
                    result['creator'] = self._extract_xml_value(meta_xml, 'dc:creator')
                    result['title'] = self._extract_xml_value(meta_xml, 'dc:title')
                    result['subject'] = self._extract_xml_value(meta_xml, 'dc:subject')
                    result['company'] = self._extract_xml_value(meta_xml, 'meta:company')

                internal_paths = []
                for name in names:
                    if name.endswith('.xml'):
                        try:
                            xml_content = zf.read(name).decode('utf-8', errors='ignore')
                            paths = self._find_internal_paths(xml_content)
                            internal_paths.extend(paths)
                        except Exception:
                            pass
                result['internal_paths'] = list(set(internal_paths))[:MAX_INTERNAL_PATHS]

                result['has_macros'] = False
                result['gps_coords'] = []
                result['embedded_fonts'] = []

        except Exception as e:
            logger.debug(f"[DOCMETA] ODT extraction failed: {e}")
            result = {'format': 'odt', 'error': str(e)}

        return result

    # =========================================================================
    # SVG EXTRACTION
    # =========================================================================

    def _extract_svg(self, content: bytes) -> dict:
        """Extract metadata from SVG files."""
        result: dict[str, Any] = {'format': 'svg'}

        try:
            svg_content = content.decode('utf-8', errors='ignore')

            for key, pattern in SVG_METADATA_PATTERNS.items():
                match = pattern.search(svg_content)
                if match:
                    result[key] = match.group(1).strip()

            viewbox_match = re.search(r'viewBox="([^"]*)"', svg_content)
            if viewbox_match:
                result['viewBox'] = viewbox_match.group(1)

            width_match = re.search(r'width="([^"]*)"', svg_content)
            height_match = re.search(r'height="([^"]*)"', svg_content)
            if width_match:
                result['width'] = width_match.group(1)
            if height_match:
                result['height'] = height_match.group(1)

            author_match = re.search(r'<svg[^>]*author="([^"]*)"', svg_content, re.IGNORECASE)
            if author_match:
                result['author'] = author_match.group(1)

            result['internal_paths'] = self._find_internal_paths(svg_content)[:MAX_INTERNAL_PATHS]

            result['has_macros'] = False
            result['gps_coords'] = []

        except Exception as e:
            logger.debug(f"[DOCMETA] SVG extraction failed: {e}")
            result = {'format': 'svg', 'error': str(e)}

        return result

    # =========================================================================
    # DXF EXTRACTION
    # =========================================================================

    def _extract_dxf(self, content: bytes) -> dict:
        """Extract metadata from DXF files (CAD drawings)."""
        result: dict[str, Any] = {'format': 'dxf'}

        try:
            dxf_content = content.decode('utf-8', errors='ignore')

            header_match = re.search(r'\[HEADER\](.*?)\[ENDTAB\]', dxf_content, re.DOTALL | re.IGNORECASE)
            if header_match:
                header = header_match.group(1)

                acadver = re.search(r'\$ACADVER\s*\n\s*1\s*\n([^\s]+)', header)
                if acadver:
                    result['autocad_version'] = acadver.group(1)

                insbase = re.search(r'\$INSBASE\s*\n\s*10\s*\n([^\s]+)\s*\n\s*20\s*\n([^\s]+)\s*\n\s*30\s*\n([^\s]+)', header)
                if insbase:
                    result['insertion_base'] = {
                        'x': float(insbase.group(1)),
                        'y': float(insbase.group(2)),
                        'z': float(insbase.group(3))
                    }

                extmin = re.search(r'\$EXTMIN\s*\n\s*10\s*\n([^\s]+)\s*\n\s*20\s*\n([^\s]+)', header)
                extmax = re.search(r'\$EXTMAX\s*\n\s*10\s*\n([^\s]+)\s*\n\s*20\s*\n([^\s]+)', header)
                if extmin and extmax:
                    result['coordinate_extents'] = {
                        'min': {'x': float(extmin.group(1)), 'y': float(extmin.group(2))},
                        'max': {'x': float(extmax.group(1)), 'y': float(extmax.group(2))}
                    }

            result['internal_paths'] = self._find_internal_paths(dxf_content)[:MAX_INTERNAL_PATHS]
            result['has_macros'] = False

        except Exception as e:
            logger.debug(f"[DOCMETA] DXF extraction failed: {e}")
            result = {'format': 'dxf', 'error': str(e)}

        return result

    # =========================================================================
    # EMAIL EXTRACTION
    # =========================================================================

    def _extract_email(self, content: bytes) -> dict:
        """Extract email headers with forensics analysis."""
        result: dict[str, Any] = {'format': 'eml'}

        try:
            email_content = content.decode('utf-8', errors='ignore')

            if '\n\n' in email_content:
                header_section, _ = email_content.split('\n\n', 1)
            else:
                header_section = email_content

            headers = {}
            for line in header_section.split('\n'):
                if ':' in line:
                    key, value = line.split(':', 1)
                    headers[key.strip().lower()] = value.strip()

            result['headers'] = dict(list(headers.items())[:MAX_EMAIL_HEADERS])

            result['received_chain'] = self._analyze_received_chain(headers.get('received', ''))

            originating_ip = X_ORIGINATING_IP_PATTERN.search(email_content)
            if originating_ip:
                result['originating_ip'] = originating_ip.group(1)

            msg_id = headers.get('message-id', '')
            domain_match = MESSAGE_ID_DOMAIN_PATTERN.search(msg_id)
            if domain_match:
                result['message_id_domain'] = domain_match.group(1)

            dkim = headers.get('dkim-signature', '')
            if dkim:
                dkim_domain = re.search(r'd=([^;\s]+)', dkim)
                if dkim_domain:
                    result['dkim_domain'] = dkim_domain.group(1)

            result['spf'] = headers.get('received-spf', None)
            result['from'] = headers.get('from', None)
            result['reply_to'] = headers.get('reply-to', None)
            result['date'] = headers.get('date', None)
            result['subject'] = headers.get('subject', None)

        except Exception as e:
            logger.debug(f"[DOCMETA] Email extraction failed: {e}")
            result = {'format': 'eml', 'error': str(e)}

        return result

    def _analyze_received_chain(self, received_header: str) -> list[dict[str, Any]]:
        """Analyze Received headers to build infrastructure chain."""
        chain = []
        try:
            if not received_header:
                return chain

            for line in received_header.split('\n')[:10]:
                line = line.strip()
                if not line:
                    continue

                entry: dict[str, Any] = {}

                from_match = re.search(r'from\s+([^\s\(]+)', line, re.IGNORECASE)
                if from_match:
                    entry['from'] = from_match.group(1)

                by_match = re.search(r'by\s+([^\s\(]+)', line, re.IGNORECASE)
                if by_match:
                    entry['by'] = by_match.group(1)

                ip_match = re.search(r'\(([^)]+\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}[^)]*)\)', line)
                if ip_match:
                    entry['ip'] = ip_match.group(1)

                with_match = re.search(r'with\s+(\S+)', line, re.IGNORECASE)
                if with_match:
                    entry['with'] = with_match.group(1)

                if entry:
                    chain.append(entry)

        except Exception as e:
            logger.debug(f"[DOCMETA] Received chain analysis failed: {e}")
        return chain[:10]

    def _extract_msg(self, content: bytes) -> dict:
        """Extract from Outlook MSG files."""
        result: dict[str, Any] = {'format': 'msg'}

        try:
            try:
                with zipfile.ZipFile(io.BytesIO(content)) as zf:
                    zf.namelist()
                    result['has_macros'] = False
            except Exception:
                content_str = content.decode('utf-8', errors='ignore')
                result['author'] = self._extract_email_field(content_str, 'author')
                result['subject'] = self._extract_email_field(content_str, 'subject')

        except Exception as e:
            logger.debug(f"[DOCMETA] MSG extraction failed: {e}")
            result = {'format': 'msg', 'error': str(e)}

        return result

    def _extract_email_field(self, content: str, field: str) -> str | None:
        """Extract field from email content."""
        patterns = [
            rf'{field}:\s*([^\n]+)',
            rf'"{field}"\s*:\s*"([^"]+)"',
        ]
        for pattern in patterns:
            match = re.search(pattern, content, re.IGNORECASE)
            if match:
                return match.group(1).strip()
        return None

    # =========================================================================
    # MACRO ANALYSIS
    # =========================================================================

    def _check_docx_macros(self, content: bytes) -> bool:
        """Check if DOCX/XLSX/PPTX contains VBA macros."""
        try:
            with zipfile.ZipFile(io.BytesIO(content)) as zf:
                names = zf.namelist()
                return any('vbaProject.bin' in n for n in names)
        except Exception:
            return False

    def _analyze_macros_olevba(self, content: bytes, file_type: str) -> dict[str, Any]:
        """
        Analyze macros using olevba for C2 URLs and suspicious API calls.
        Returns analysis results with threat indicators.
        """
        analysis: dict[str, Any] = {
            'c2_urls': [],
            'suspicious_api_calls': [],
            'macro_strings': [],
            'is_suspicious': False,
            'threat_level': 'low',
        }

        if not OLEVB_AVAILABLE:
            return analysis

        try:
            vba_code = self._extract_vba_code(content)
            if not vba_code:
                return analysis

            for pattern in C2_URL_PATTERNS:
                for match in pattern.finditer(vba_code):
                    url = match.group(0)
                    if url not in analysis['c2_urls']:
                        analysis['c2_urls'].append(url)
                        analysis['is_suspicious'] = True

            for pattern in SUSPICIOUS_API_PATTERNS:
                for match in pattern.finditer(vba_code):
                    api_call = match.group(0)
                    if api_call not in analysis['suspicious_api_calls']:
                        analysis['suspicious_api_calls'].append(api_call)
                        analysis['is_suspicious'] = True

            string_pattern = re.compile(r'"([^"]{4,100})"')
            for match in string_pattern.finditer(vba_code):
                s = match.group(1)
                if any(c.isdigit() for c in s) and len(s) > 8:
                    analysis['macro_strings'].append(s)

            if len(analysis['c2_urls']) > 0 or len(analysis['suspicious_api_calls']) > 3:
                analysis['threat_level'] = 'high'
            elif len(analysis['suspicious_api_calls']) > 0:
                analysis['threat_level'] = 'medium'

        except Exception as e:
            logger.debug(f"[DOCMETA] Macro analysis failed: {e}")

        return analysis

    def _extract_vba_code(self, content: bytes) -> str:
        """Extract VBA code from Office documents."""
        vba_code = []
        try:
            with zipfile.ZipFile(io.BytesIO(content)) as zf:
                for name in zf.namelist():
                    if 'vbaProject.bin' in name or name.endswith('.cls') or name.endswith('.bas'):
                        try:
                            vba_data = zf.read(name)
                            try:
                                vba_code.append(vba_data.decode('utf-8', errors='ignore'))
                            except Exception:
                                vba_code.append(vba_data.decode('latin-1', errors='ignore'))
                        except Exception:
                            pass
        except Exception:
            pass
        return '\n'.join(vba_code)[:MAX_MACRO_ANALYSIS_CHARS]

    # =========================================================================
    # UTILITIES
    # =========================================================================

    def _extract_xml_value(self, xml_content: str, tag: str) -> str | None:
        """Extract value from XML tag."""
        patterns = [
            rf'<{tag}>([^<]+)</{tag}>',
            rf'<{tag}\s+[^>]*>([^<]+)</{tag}>',
        ]
        for pattern in patterns:
            match = re.search(pattern, xml_content, re.IGNORECASE)
            if match:
                return match.group(1).strip()
        return None

    def _find_internal_paths(self, text: str) -> list[str]:
        """Find internal file paths in text."""
        paths = set()
        for _pattern_type, pattern in INTERNAL_PATH_PATTERNS.items():
            matches = pattern.findall(text)
            paths.update(matches)
        return list(paths)


def _exif_to_float(val):
    """Handle EXIF rational (num, denom) tuples and plain numeric values."""
    if isinstance(val, (tuple, list)):
        try:
            return val[0] / val[1]
        except (ZeroDivisionError, TypeError):
            return 0.0
    return float(val) if val else 0.0
